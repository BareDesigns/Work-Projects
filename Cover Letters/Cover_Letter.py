import openpyxl
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Setting Variable Names and lists
document = Document("G:\Enrollment Management Center\Graduation\default.docx")
os.chdir("G:\Enrollment Management Center\Graduation")
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(20)
fName = []
lName = []
gNum = []
pCode = []

# Getting the file name
fileName = input('Drag the file you want to use here:\n')
wb = openpyxl.load_workbook(fileName)
sheet = wb.get_sheet_by_name('Sheet1')


# Centering the text
def indent(word):
    word_format = word.paragraph_format
    word.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return word_format


# Text formatting for title lines
def header(para):
    para.bold = True
    para.underline = True
    para.font.size = Pt(25)
    return para


# Add values to list
for i in range(1, sheet.max_row+1):
    fName.append(sheet.cell(row=i, column=6).value)
    lName.append(sheet.cell(row=i, column=7).value)
    gNum.append(sheet.cell(row=i, column=8).value)
    pCode.append(sheet.cell(row=i, column=9).value)

# Writing the lists to Word
for item in range(len(fName)):

    pOne = document.add_paragraph()
    p1 = pOne.add_run("\n\n\n\nSTUDENT'S NAME:")
    indent(pOne)
    header(p1)

    pTwo = document.add_paragraph(fName[item] + ' ' + lName[item] + '\n')
    indent(pTwo)

    pThree = document.add_paragraph()
    p3 = pThree.add_run("\nG NUMBER:")
    indent(pThree)
    header(p3)

    pFour = document.add_paragraph(gNum[item])
    indent(pFour)

    pFive = document.add_paragraph()
    p5 = pFive.add_run("\nProgram Codes:")
    indent(pFive)
    header(p5)

    pSix = document.add_paragraph(pCode[item])
    indent(pSix)

    document.add_page_break()

saveFile = input('What do you want to name the file?\n')
document.save(saveFile + '.docx')
print('\nMaking New Spreadsheet...')
time.sleep(4)
exit = input('\n COMPLETE! Press ENTER to close the program')
