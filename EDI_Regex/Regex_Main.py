import re
import time
import os
from docx import Document

document = Document("G:\Enrollment Management Center\Graduation\default.docx")
fileLocation = input('Drop file here (Delete quotations around file name):\n')
os.chdir("G:\Enrollment Management Center\Evaluation Requests\EDIs")


class CleanLine:
    '''This class cleans the extra blank spaces and changes the line according
    to the data inside'''

    def __init__(self, line):
        '''Removes all the extra blank spaces from everyline of the EDI'''
        self.line = re.sub(r'(\s+)', ' ', line.strip())

    def course(self):
        '''This function removes all of the unnecessary information from the
        lines containing coursework'''
        self.line = re.sub(r'(\W)$', '', self.line)
        self.line = re.sub(r'((\s)(\d))$', '', self.line)
        self.line = re.sub(r'((\s)([E|I]))$', '', self.line)
        self.line = re.sub(r'((I/F))$', 'F', self.line)
        self.line = re.sub(r'((WF|WQ|FX))$', 'F', self.line)
        self.line = re.sub(r'((WL|WP))$', 'W', self.line)
        self.line = re.sub(r'^(WBCT).+', '', self.line)
        self.line = re.sub(r'(ZZZ)$', 'IP', self.line)
        self.line = re.sub(r'(#)(.*)$', 'IP', self.line)
        return self.line

    def names(self):
        '''Makes the line containing names prettier'''
        self.line = re.sub(r'AS: ', '', self.line)
        return self.line

    def terms(self):
        '''Changes the term line to a simple *TERM* *YEAR* '''
        if bool(re.search(r'(Mini:)', self.line, re.IGNORECASE)) is True:
            self.line = re.sub(r'\s*(<<).+', 'MINIMESTER', self.line)
            return self.line

        elif bool(re.search(r'(Correspondence)',
                            self.line, re.IGNORECASE)) is True:
            self.line = re.sub(r'\s*(<<).+', 'CORRESPONDENCE', self.line)
            return self.line

        elif bool(re.search(r'(Quarter)',
                            self.line, re.IGNORECASE)) is True:
            self.line = re.sub(r'\s*(<<).+', 'Quarter', self.line)
            return self.line

        elif bool(re.search(r'(Orientation)', self.line,
                            re.IGNORECASE)) is True:
            self.line = re.sub(r'\s*(<<).+', 'ORIENTATION', self.line)
            return self.line

        elif bool(re.search(r'^\s*(<<:|<<Non)', self.line)) is True:
            self.line = re.sub(r'\s*(<<:).+',
                               '??LIKELY TRNSFRWRK??', self.line)
            return self.line

        else:
            mo = re.compile(r'(spr|sum|fall|spring|winter|may).+?(\d{4})',
                            re.IGNORECASE)
            foo = mo.search(self.line)
            term, year = foo.groups()
            self.line = re.sub(r'(\s*).+', '' + term.upper() + ' ' + year, self.line)
            return self.line


with open(fileLocation, 'r') as transcript:
    data = transcript.readlines()
    for line in data:

        new_line = CleanLine(line)

        if new_line.line.startswith('G0'):
            p = document.add_paragraph('= ' * 50 + '\n')
            p.add_run('\n' + new_line.line + '\n').bold = True

        elif re.match(r'^((\d){6})', line):
            p.add_run(new_line.line + '\n').bold = True

        elif new_line.line.startswith('AS: '):
            p.add_run('\n' + new_line.names())

        elif new_line.line.startswith('Student Name: '):
            p.add_run('\n' + new_line.names() + '\n').italic = True

        elif new_line.line.startswith('<<'):
            p.add_run('\n' + new_line.terms() + '\n')

        elif re.search(r'^([A-Z]{1,10})(\s)(\d)', new_line.line):
            p.add_run(new_line.course() + '\n')

        elif re.search(r'^([A-Z]{1})(\s)([A-Z]{1,4})(\s)(\d){1,7}',
                       new_line.line):
            p.add_run(new_line.course() + '\n')

        elif new_line.line.startswith('* Inst Qual:'):
            p.add_run('!!!!!!!!!!!!!!TRANSFER HOURS!!!!!!!!!!!!!!' +
                      '\n').italic = True

        else:
            pass

fileName = input('\nWhat would you like to name the new file?:\n')
document.save(fileName + '.docx')
print("\nSaving File...")
# Add loading time for peace of mind
time.sleep(3)
exit = input("Press ENTER to close the program")
